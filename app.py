# from flask import Flask, render_template, request, redirect, url_for, flash
# from werkzeug.utils import secure_filename
# import os
# import spacy
# import pandas as pd
# from pdfminer.high_level import extract_text
# from docx import Document
# from flask import request, jsonify

# @app.route('/process-file', methods=['POST'])


#     # Your existing logic to extract text and skills
#     # ...


# app = Flask(__name__)
# app.secret_key = 'your_secret_key'  # Needed for flashing messages
# UPLOAD_FOLDER = 'uploads'
# ALLOWED_EXTENSIONS = {'pdf', 'docx'}
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# nlp = spacy.load("en_core_web_sm")
# skills_df = pd.read_csv('skills.csv')
# skills_list = skills_df.columns.tolist()

# if not os.path.exists(app.config['UPLOAD_FOLDER']):
#     os.makedirs(app.config['UPLOAD_FOLDER'])

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(pdf_path):
#     return extract_text(pdf_path)

# def extract_text_from_docx(docx_path):
#     doc = Document(docx_path)
#     return " ".join(paragraph.text for paragraph in doc.paragraphs)

# def extract_skills(text, skills_list):
#     doc = nlp(text)
#     found_skills = set()
#     for skill in skills_list:
#         if skill.lower() in text.lower():
#             found_skills.add(skill)
#     return list(found_skills)

# @app.route('/', methods=['GET', 'POST'])




# os.remove(file_path)
# return jsonify({'skills': extracted_skills})

# if __name__ == '__main__':
#     app.run(debug=True)

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
from werkzeug.utils import secure_filename
import os
import spacy
import pandas as pd
from pdfminer.high_level import extract_text
from docx import Document

app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Replace with a real secret key

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

nlp = spacy.load("en_core_web_sm")
skills_df = pd.read_csv('skills.csv')
skills_list = skills_df.columns.tolist()

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    return extract_text(pdf_path)

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return " ".join(paragraph.text for paragraph in doc.paragraphs)

def extract_skills(text, skills_list):
    doc = nlp(text)
    found_skills = set()
    for skill in skills_list:
        if skill.lower() in text.lower():
            found_skills.add(skill)
    return list(found_skills)

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'resume' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['resume']
        if file.filename == '' or not allowed_file(file.filename):
            flash('No selected file or file type not allowed')
            return redirect(request.url)
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            file_type = 'pdf' if filename.endswith('.pdf') else 'docx'
            text = extract_text_from_pdf(file_path) if file_type == 'pdf' else extract_text_from_docx(file_path)
            extracted_skills = extract_skills(text, skills_list)
        except Exception as e:
            flash('Error processing file')
            return redirect(request.url)
        finally:
            os.remove(file_path)

        return jsonify({'skills': extracted_skills})
    return render_template('upload.html')

if __name__ == '__main__':
    app.run(debug=True)
